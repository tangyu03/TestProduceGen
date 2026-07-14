#!/usr/bin/env python3
"""Extract full text content from the requirements docx file."""
from docx import Document
import sys

doc_path = "/home/z/my-project/upload/网数中心能力验证服务平台升级维护项目-歧义修正.docx"
doc = Document(doc_path)

output_lines = []

# Extract paragraphs with style info
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text:
        style = para.style.name if para.style else "Normal"
        output_lines.append(f"[P{i:04d}][{style}] {text}")

# Extract tables
for ti, table in enumerate(doc.tables):
    output_lines.append(f"\n===== TABLE {ti} =====")
    for ri, row in enumerate(table.rows):
        cells = [cell.text.strip().replace("\n", " | ") for cell in row.cells]
        output_lines.append(f"  ROW{ri}: " + " || ".join(cells))
    output_lines.append(f"===== END TABLE {ti} =====\n")

full_text = "\n".join(output_lines)

# Write to file
with open("/home/z/my-project/scripts/doc_content.txt", "w", encoding="utf-8") as f:
    f.write(full_text)

print(f"Total paragraphs: {len(doc.paragraphs)}")
print(f"Total tables: {len(doc.tables)}")
print(f"Total chars: {len(full_text)}")
print("Saved to /home/z/my-project/scripts/doc_content.txt")
